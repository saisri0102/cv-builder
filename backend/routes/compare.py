# backend/routes/compare.py
from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi import status
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any
import os
import tempfile
import logging
import unicodedata
import re
import time
import asyncio

import pdfplumber
import docx

# Use the improved engine you have in backend/utils/similarity.py
from backend.utils.similarity import compute_similarity

router = APIRouter(prefix="/compare", tags=["Compare"])

# ---- settings / limits ----
ALLOWED_EXTS = {".pdf", ".docx", ".txt"}  # .doc is flaky; keep docx/txt/pdf
ALLOWED_MIMES = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "text/plain": ".txt",
}
MAX_BYTES = 5 * 1024 * 1024  # 5 MB
DEFAULT_MAX_PDF_PAGES = 40   # hard cap for huge PDFs

# ---------- text cleanup helpers ----------
_dehyphen_re = re.compile(r"(\w)-\s*\n\s*(\w)")
_multispace_re = re.compile(r"[ \t\f\v]+")


def _normalize_text(s: str) -> str:
    """
    Normalize Unicode, remove hyphenation across line breaks, collapse whitespace,
    and lowercase. Keeps content as plain text for the matcher.
    """
    if not s:
        return ""
    s = unicodedata.normalize("NFC", s)
    s = _dehyphen_re.sub(r"\1 \2", s)
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = "\n".join(_multispace_re.sub(" ", line).strip() for line in s.split("\n"))
    return s.lower().strip()


# ---------- file helpers ----------
def _save_to_temp(upload: UploadFile) -> str:
    """Write UploadFile stream to a temp file and return the path."""
    upload.file.seek(0)
    fd, path = tempfile.mkstemp()
    total = 0
    with os.fdopen(fd, "wb") as out:
        for chunk in iter(lambda: upload.file.read(1024 * 1024), b""):
            total += len(chunk)
            if total > MAX_BYTES:
                # Write nothing further; delete below
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"File {upload.filename} exceeds 5MB limit",
                )
            out.write(chunk)
    return path


def _validate_ext_and_mime(path: str, filename: str, content_type: Optional[str]):
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXTS:
        raise HTTPException(status_code=400, detail=f"Unsupported file format: {ext}")

    # MIME check is best-effort (browsers sometimes send 'application/octet-stream')
    if content_type and content_type in ALLOWED_MIMES:
        expected_ext = ALLOWED_MIMES[content_type]
        if expected_ext != ext:
            logging.warning(
                f"[compare] MIME/ext mismatch for {filename}: {content_type} vs {ext}"
            )

    size = os.path.getsize(path)
    if size > MAX_BYTES:
        raise HTTPException(status_code=400, detail=f"File {filename} exceeds 5MB limit")


def _extract_pdf_text(temp_path: str, max_pages: int) -> str:
    with pdfplumber.open(temp_path) as pdf:
        parts = []
        page_count = min(len(pdf.pages), max_pages)
        for i in range(page_count):
            try:
                page = pdf.pages[i]
                txt = page.extract_text() or ""
                parts.append(txt)
            except Exception as e:
                logging.warning(f"[compare] PDF page {i} extraction failed: {e}")
                parts.append("")
        if len(pdf.pages) > max_pages:
            logging.info(f"[compare] Truncated PDF to first {max_pages} pages")
        return "\n".join(parts)


def _extract_docx_text(temp_path: str) -> str:
    d = docx.Document(temp_path)
    lines = [(p.text or "") for p in d.paragraphs]
    # Also collect table cell text (common for resumes exported from templates)
    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text:
                    lines.append(cell.text)
    return "\n".join(lines)


def _extract_txt_text(temp_path: str) -> str:
    # Try utf-8, then latin-1 as permissive fallback
    try:
        with open(temp_path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        with open(temp_path, "r", encoding="latin-1", errors="ignore") as f:
            return f.read()


def extract_text(upload: UploadFile, *, max_pdf_pages: int = DEFAULT_MAX_PDF_PAGES) -> str:
    """Extract plain text from PDF, DOCX, or TXT. Returns normalized, lowercased text."""
    temp_path = _save_to_temp(upload)
    try:
        _validate_ext_and_mime(temp_path, upload.filename, getattr(upload, "content_type", None))
        ext = os.path.splitext(upload.filename)[1].lower()

        if ext == ".pdf":
            text = _extract_pdf_text(temp_path, max_pages=max_pdf_pages)
        elif ext == ".docx":
            text = _extract_docx_text(temp_path)
        elif ext == ".txt":
            text = _extract_txt_text(temp_path)
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file: {ext}")

        return _normalize_text(text)
    except HTTPException:
        raise
    except Exception as e:
        logging.exception(f"[compare] Failed to extract text from {upload.filename}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to extract text from {upload.filename}")
    finally:
        try:
            os.remove(temp_path)
        except Exception:
            pass


# ---------- response models ----------
class ResumeComparison(BaseModel):
    fileName: str = Field(..., description="Original file name")
    match_percentage: Optional[float] = Field(None, description="Similarity score (0-100)")
    matched_keywords: Optional[List[str]] = None
    unmatched_keywords: Optional[List[str]] = None
    resume_text: Optional[str] = None
    error: Optional[str] = None


class CompareResponse(BaseModel):
    jd_text: Optional[str] = None
    results: List[ResumeComparison]
    summary: Dict[str, Any]


# ---------- worker ----------
async def _process_single_resume(
    upload: UploadFile,
    jd_content: str,
    comparison_type: str,
    *,
    return_text: bool,
    max_pdf_pages: int,
    top_n_keywords: Optional[int],
) -> ResumeComparison:
    try:
        # Extraction can be CPU/IO heavy; push to thread
        resume_text = await asyncio.to_thread(
            extract_text, upload, max_pdf_pages=max_pdf_pages
        )
        if not resume_text.strip():
            return ResumeComparison(
                fileName=upload.filename,
                error="Resume file contains no extractable text",
                resume_text=resume_text if return_text else None,
            )

        sim_kwargs = {}
        if top_n_keywords is not None:
            sim_kwargs["top_n_keywords"] = top_n_keywords

        sim = await asyncio.to_thread(
            compute_similarity, resume_text, jd_content, comparison_type, **sim_kwargs
        )
        return ResumeComparison(
            fileName=upload.filename,
            match_percentage=sim.get("match_percentage"),
            matched_keywords=sim.get("matched_keywords"),
            unmatched_keywords=sim.get("unmatched_keywords"),
            resume_text=resume_text if return_text else None,
        )
    except HTTPException as http_err:
        return ResumeComparison(fileName=upload.filename, error=http_err.detail)
    except Exception as e:
        logging.exception(f"[compare] Unexpected error with {upload.filename}: {e}")
        return ResumeComparison(fileName=upload.filename, error="Unexpected error while processing")


# ---------- route ----------
@router.post("") 
async def compare_resumes(
    resumes: List[UploadFile] = File(..., description="One or more resumes (.pdf, .docx, .txt)"),
    jd_file: Optional[UploadFile] = File(None, description="Job description file (.pdf, .docx, .txt)"),
    jd_text: Optional[str] = Form("", description="Raw job description text"),
    comparison_type: Optional[str] = Form("word"),
    # New optional knobs:
    return_text: bool = Form(True, description="Include normalized texts in response"),
    max_pdf_pages: int = Form(DEFAULT_MAX_PDF_PAGES, description="Truncate PDF parsing to this many pages"),
    top_n_keywords: Optional[int] = Form(None, description="If provided, limit matched/unmatched lists"),
):
    """
    Compare each uploaded resume to the job description (file or text).

    Modes from utils.similarity:
      - word: smarter token/phrase overlap (fuzzy tolerant)
      - skill: overlap within curated skills set (plus fuzzy)
      - overall: union of word/skill/phrase matches

    Extras:
      - return_text: include resume/jd text in the response (useful for subsequent enhancement steps)
      - max_pdf_pages: safety cap for very large PDFs
      - top_n_keywords: optional limit for keyword lists returned by compute_similarity
    """
    t0 = time.time()
    if comparison_type not in {"word", "skill", "overall"}:
        raise HTTPException(status_code=400, detail=f"Invalid comparison type: {comparison_type}")

    if not resumes:
        raise HTTPException(status_code=400, detail="No resumes uploaded")

    # JD content (file overrides text if both are supplied)
    if jd_file is not None:
        jd_content = await asyncio.to_thread(
            extract_text, jd_file, max_pdf_pages=max_pdf_pages
        )
        jd_content = jd_content.strip()
    else:
        jd_content = _normalize_text(jd_text or "").strip()

    if not jd_content:
        raise HTTPException(status_code=400, detail="Job description is empty")

    # Process resumes concurrently
    tasks = [
        _process_single_resume(
            upload=r,
            jd_content=jd_content,
            comparison_type=comparison_type,
            return_text=return_text,
            max_pdf_pages=max_pdf_pages,
            top_n_keywords=top_n_keywords,
        )
        for r in resumes
    ]
    results: List[ResumeComparison] = await asyncio.gather(*tasks)

    # Compute summary stats
    scored = [r for r in results if r.match_percentage is not None]
    best_match = None
    avg_score = None
    if scored:
        best_match = max(scored, key=lambda x: (x.match_percentage or 0))
        avg_score = sum((x.match_percentage or 0.0) for x in scored) / len(scored)

    elapsed_ms = int((time.time() - t0) * 1000)
    summary: Dict[str, Any] = {
        "count": len(results),
        "processed_ms": elapsed_ms,
        "average_match": round(avg_score, 2) if avg_score is not None else None,
        "best_match": {
            "fileName": best_match.fileName,
            "match_percentage": best_match.match_percentage,
        } if best_match else None,
        "comparison_type": comparison_type,
        "truncated_pdf_pages": max_pdf_pages,
    }

    return CompareResponse(
        jd_text=jd_content if return_text else None,
        results=results,
        summary=summary,
    )
